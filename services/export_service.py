"""
Export Service
Generate reports in various formats
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config.settings import Settings
from config.database import get_db
from database.crud.projects import get_project
from database.crud.stages import get_all_stage_progress
from datetime import datetime
import json

def generate_final_report(project_id: int, format: str = 'docx', sections: list = None) -> str:
    """
    Generate comprehensive final report

    Args:
        project_id: Project ID
        format: Report format (pdf, docx)
        sections: List of sections to include

    Returns:
        Path to generated report
    """
    db = get_db()
    project = get_project(db, project_id)
    stages = get_all_stage_progress(db, project_id)

    if format == 'docx' or format == 'both':
        report_path = _generate_docx_report(project, stages, sections)
    else:
        report_path = _generate_docx_report(project, stages, sections)

    db.close()
    return report_path

def _generate_docx_report(project, stages, sections) -> str:
    """Generate DOCX report"""
    doc = Document()

    # Title
    title = doc.add_heading('Design Thinking Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project Information
    doc.add_heading('Project Information', 1)
    doc.add_paragraph(f'Project Name: {project.name}')
    doc.add_paragraph(f'Area: {project.area}')
    doc.add_paragraph(f'Goal: {project.goal}')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d")}')

    doc.add_page_break()

    # Executive Summary
    if 'Executive Summary' in sections:
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(
            f'This report documents the Design Thinking process for {project.name} '
            f'in the {project.area} domain. The project aimed to {project.goal}.'
        )
        doc.add_paragraph('')

    # Stage sections
    stage_names = ['Empathise', 'Define', 'Ideate', 'Prototype', 'Test', 'Implement']

    for stage in stages:
        stage_name = stage_names[stage.stage_number - 1]

        if stage_name in sections:
            doc.add_heading(f'Stage {stage.stage_number}: {stage_name}', 1)
            doc.add_paragraph(f'Status: {stage.status.replace("_", " ").title()}')

            if stage.data:
                doc.add_heading('Stage Data:', 2)
                for key, value in stage.data.items():
                    doc.add_paragraph(f'{key.replace("_", " ").title()}: {value}', style='List Bullet')

            if stage.completed_at:
                doc.add_paragraph(f'Completed: {stage.completed_at.strftime("%Y-%m-%d")}')

            doc.add_page_break()

    # Conclusion
    doc.add_heading('Conclusion', 1)
    completed_stages = sum(1 for s in stages if s.status == 'completed')
    doc.add_paragraph(
        f'This Design Thinking project has completed {completed_stages} out of 6 stages. '
        f'The insights and outputs generated through this structured process provide a solid '
        f'foundation for implementation and future iterations.'
    )

    # Save document
    report_path = Settings.EXPORT_DIR / f'report_{project.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(report_path)

    return str(report_path)

def export_stage_data(project_id: int, stage_number: int, format: str = 'json') -> str:
    """
    Export specific stage data

    Args:
        project_id: Project ID
        stage_number: Stage number
        format: Export format (json, csv)

    Returns:
        Path to exported file
    """
    from database.crud.stages import get_stage_progress

    db = get_db()
    stage = get_stage_progress(db, project_id, stage_number)
    db.close()

    if not stage:
        raise ValueError(f"Stage {stage_number} not found for project {project_id}")

    export_path = Settings.EXPORT_DIR / f'stage_{stage_number}_{project_id}.{format}'

    if format == 'json':
        with open(export_path, 'w') as f:
            json.dump(stage.data, f, indent=2)

    return str(export_path)
